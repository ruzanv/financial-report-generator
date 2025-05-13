import os
import logging
import tempfile
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import skew, kurtosis
from statsmodels.graphics.tsaplots import plot_acf
from docx import Document
from docx.shared import Inches

logger = logging.getLogger('docx_generator')

def create_report(df, preds, report_path):

    logger.info('Building report', extra={'path': str(report_path)})

    definitions = {
        'line_1600': 'Выручка (Revenue) — итоговые поступления от продаж.',
        'line_2110': 'Себестоимость (Cost of Goods Sold) — затраты на производство товаров.',
        'line_2120': 'Операционная прибыль (Operating Income) — прибыль до учёта финансовых расходов.',
        'line_2400': 'Денежный поток (Cash Flow) — чистый приток/отток денежных средств.'
    }

    labels = {
        'line_1600': 'Выручка',
        'line_2110': 'Себестоимость',
        'line_2120': 'Операционная прибыль',
        'line_2400': 'Денежный поток'
    }

    xgb_preds = np.array(preds.get('xgboost', []), dtype=float)
    lstm_preds = np.array(preds.get('lstm', []), dtype=float)
    xgb_model = preds.get('xgboost_model', None)

    doc = Document()
    doc.add_heading('Отчёт о финансовом прогнозе', level=1)
    doc.add_paragraph(
        'Спасибо за то, что выбрали наш продукт. '
        'В этом отчёте представлены прогнозы ключевых финансовых показателей на следующий период и их интерпретация.'
    )

    doc.add_heading('Расшифровка показателей', level=2)
    for desc in definitions.values():
        doc.add_paragraph(desc, style='List Number')

    doc.add_heading('Исходные данные от вас', level=2)
    tbl = doc.add_table(rows=len(df) + 1, cols=len(labels))
    tbl.style = 'Table Grid'
    for j, col in enumerate(labels):
        tbl.cell(0, j).text = labels[col]
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, col in enumerate(labels):
            val = getattr(row, col)
            tbl.cell(i, j).text = f'{val:,.2f}' if isinstance(val, float) else str(val)

    doc.add_heading('Прогнозы по моделям', level=2)
    for model_name, arr in [('XGBoost', xgb_preds), ('LSTM', lstm_preds)]:
        doc.add_heading(model_name, level=3)
        tblp = doc.add_table(rows=len(arr) + 1, cols=2)
        tblp.style = 'Table Grid'
        tblp.cell(0, 0).text, tblp.cell(0, 1).text = '№', 'Значение'
        for i, v in enumerate(arr, start=1):
            tblp.cell(i, 0).text = str(i)
            tblp.cell(i, 1).text = f'{v:.2f}'

    doc.add_heading('Ключевые статистические метрики прогнозов', level=2)
    def add_basic_stats(title, arr):
        doc.add_heading(title, level=3)
        stats = {
            'Среднее': np.mean(arr),
            'Медиана': np.median(arr),
            'Стандартное отклонение': np.std(arr),
            'Минимум': np.min(arr),
            'Максимум': np.max(arr)
        }
        tbls = doc.add_table(rows=len(stats), cols=2)
        tbls.style = 'Table Grid'
        for i, (n, v) in enumerate(stats.items()):
            tbls.cell(i,0).text, tbls.cell(i,1).text = n, f'{v:.2f}'
    add_basic_stats('XGBoost', xgb_preds)
    add_basic_stats('LSTM', lstm_preds)

    doc.add_heading('Базовые статистические метрики прогнозов', level=2)

    def add_basic_stats(title, arr):
        doc.add_heading(title, level=3)
        stats = {
            'Среднее': np.mean(arr),
            'Медиана': np.median(arr),
            'Стандартное отклонение': np.std(arr),
            'Минимум': np.min(arr),
            'Максимум': np.max(arr)
        }
        tbl = doc.add_table(rows=len(stats), cols=2)
        tbl.style = 'Table Grid'
        for i, (n, v) in enumerate(stats.items()):
            tbl.cell(i, 0).text = n
            tbl.cell(i, 1).text = f'{v:.2f}'

    doc.add_paragraph(
        'Среднее — арифметическая средняя прогнозных значений, показывает общий уровень прогноза; ' \
        'Медиана — центральное значение, устойчива к выбросам; ' \
        'Стандартное отклонение — мера разброса относительно среднего; ' \
        'Минимум и Максимум — крайние границы диапазона прогноза.',
        style='Intense Quote'
    )

    add_basic_stats('XGBoost — основные метрики', xgb_preds)
    add_basic_stats('LSTM — основные метрики', lstm_preds)

    doc.add_heading('Расширенные описательные статистики', level=2)

    def add_extended_stats(title, arr):
        doc.add_heading(title, level=3)
        stats = {
            'Коэффициент вариации (CV = Std/Mean)': (np.std(arr) / np.mean(arr) if np.mean(arr) != 0 else np.nan),
            'Скос (Skewness)': skew(arr),
            'Эксцесс (Kurtosis)': kurtosis(arr),
            '25-й квантиль (Q1)': np.percentile(arr, 25),
            '75-й квантиль (Q3)': np.percentile(arr, 75),
            'Интерквантильный размах (IQR)': np.percentile(arr, 75) - np.percentile(arr, 25)
        }
        tbl = doc.add_table(rows=len(stats), cols=2)
        tbl.style = 'Table Grid'
        for i, (n, v) in enumerate(stats.items()):
            tbl.cell(i, 0).text = n
            tbl.cell(i, 1).text = f'{v:.2f}'

    doc.add_paragraph(
        'CV — относительная волатильность прогноза; Скос — асимметрия распределения; ' \
        'Эксцесс — степень «пиковости» или «плоскости» по сравнению с нормальным; ' \
        'Квантильные метрики (Q1, Q3, IQR) характеризуют расположение и размах середины распределения.',
        style='Intense Quote'
    )

    add_extended_stats('XGBoost — расширенные метрики', xgb_preds)
    add_extended_stats('LSTM — расширенные метрики', lstm_preds)

    # дополнительные вычисляемые метрики
    doc.add_heading('Дополнительные вычисляемые метрики', level=2)

    def add_additional_metrics(title, arr):
        doc.add_heading(title, level=3)
        mad = np.mean(np.abs(arr - np.mean(arr)))
        data_range = np.ptp(arr)
        slope = np.polyfit(np.arange(len(arr)), arr, 1)[0]
        stats = {
            'Среднее абсолютное отклонение (MAD)': mad,
            'Размах (max-min)': data_range,
            'Коэффициент тренда (Slope)': slope
        }
        tbl = doc.add_table(rows=len(stats), cols=2)
        tbl.style = 'Table Grid'
        for i, (n, v) in enumerate(stats.items()):
            tbl.cell(i, 0).text = n
            tbl.cell(i, 1).text = f'{v:.4f}'
    doc.add_paragraph(
        'MAD — средняя величина отклонений от среднего, показывает типичную амплитуду колебаний; ' \
        'Размах — полная амплитуда прогноза; ' \
        'Коэффициент тренда — средний прирост (отклонение) прогноза за единицу периода.',
        style='Intense Quote'
    )

    add_additional_metrics('XGBoost — вычисляемые метрики', xgb_preds)
    add_additional_metrics('LSTM — вычисляемые метрики', lstm_preds)

    # динамика прогнозов
    fig, ax = plt.subplots()
    ax.plot(xgb_preds, label='XGBoost')
    ax.plot(lstm_preds, '--', label='LSTM')
    ax.legend(); ax.set_title('Динамика прогнозов')
    ax.set_xlabel('Период'); ax.set_ylabel('Значение')
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        fig.savefig(tmp.name); dyn = tmp.name
    plt.close(fig)
    doc.add_picture(dyn, width=Inches(6))
    doc.add_heading('Динамика прогнозов по периодам', level=3)
    doc.add_paragraph('Отображает изменение прогнозных значений во времени для обеих моделей.', style='Intense Quote')

    # скользящее среднее и волатильность
    def add_rolling(doc, title, arr, window=5):
        roll_mean = pd.Series(arr).rolling(window).mean()
        roll_std  = pd.Series(arr).rolling(window).std()
        fig, ax = plt.subplots()
        ax.plot(roll_mean, label='Скользящее среднее')
        ax.fill_between(range(len(roll_std)), roll_mean-roll_std, roll_mean+roll_std, alpha=0.3)
        ax.set_title('Скользящее среднее ± σ (' + title + ')')
        ax.set_xlabel('Период'); ax.set_ylabel('Значение')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); rimg = tmp.name
        plt.close(fig)
        doc.add_picture(rimg, width=Inches(6))
        doc.add_heading('Скользящее среднее и волатильность — ' + title, level=3)
        doc.add_paragraph('Полоса ±σ показывает периоды повышенной и пониженной волатильности прогноза.',	style='Intense Quote')
    add_rolling(doc, 'XGBoost', xgb_preds)
    add_rolling(doc, 'LSTM', lstm_preds)

    # автокорреляция
    def add_acf(doc, title, arr):
        fig, ax = plt.subplots()
        plot_acf(arr, lags=20, ax=ax)
        ax.set_title('Автокорреляция прогнозов (' + title + ')')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); aimg = tmp.name
        plt.close(fig)
        doc.add_picture(aimg, width=Inches(6))
        doc.add_heading('Автокорреляция — ' + title, level=3)
        doc.add_paragraph('Пики на лаге k указывают на связь прогнозов, отстоящих на k периодов.', style='Intense Quote')
    add_acf(doc, 'XGBoost', xgb_preds)
    add_acf(doc, 'LSTM', lstm_preds)

    # boxplot, violin, KDE
    def add_distribution(doc, title, arr):
        # Boxplot
        fig, ax = plt.subplots()
        ax.boxplot(arr, vert=False)
        ax.set_title('Box-plot прогнозных значений (' + title + ')')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); bimg = tmp.name
        plt.close(fig)
        doc.add_picture(bimg, width=Inches(6))
        doc.add_heading('Box-plot — ' + title, level=3)
        doc.add_paragraph('Показывает медиану, IQR и выбросы распределения.', style='Intense Quote')

        fig, ax = plt.subplots()
        ax.violinplot(arr, vert=False)
        ax.set_title('Violin-plot прогнозных значений (' + title + ')')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); vimg = tmp.name
        plt.close(fig)
        doc.add_picture(vimg, width=Inches(6))
        doc.add_heading('Violin-plot — ' + title, level=3)
        doc.add_paragraph('Отображает форму распределения и плотность значений.', style='Intense Quote')
        fig, ax = plt.subplots()
        sns.kdeplot(arr, ax=ax, fill=True)
        ax.set_title('Оценка плотности (KDE) — ' + title)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); kimg = tmp.name
        plt.close(fig)
        doc.add_picture(kimg, width=Inches(6))
        doc.add_heading('Оценка плотности (KDE) — ' + title, level=3)
        doc.add_paragraph('Гладкая кривая плотности подчёркивает моды и формы распределения.', style='Intense Quote')
    add_distribution(doc, 'XGBoost', xgb_preds)
    add_distribution(doc, 'LSTM', lstm_preds)

    def add_pred_corr(doc, arr1, arr2):
        # pd.Series чтобы учесть разную длину
        preds_df = pd.DataFrame({
            'XGBoost': pd.Series(arr1),
            'LSTM':    pd.Series(arr2)
        })
        corr = preds_df.corr()
        fig, ax = plt.subplots()
        sns.heatmap(corr, annot=True, fmt='.2f', cmap='coolwarm', ax=ax)
        ax.set_title('Корреляция прогнозов моделей')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name); cimg = tmp.name
        plt.close(fig)
        doc.add_picture(cimg, width=Inches(6))
        doc.add_heading('Корреляция прогнозов моделей', level=3)
        doc.add_paragraph('Показывает степень линейной связи между прогнозами двух моделей. ' 
                          'Используем pd.Series, чтобы учесть различную длину массивов.',
                          style='Intense Quote')
    add_pred_corr(doc, xgb_preds, lstm_preds)

    if xgb_model is not None and hasattr(xgb_model, 'feature_importances_'):
        fi = xgb_model.feature_importances_
        feats = list(labels.values())
        fig, ax = plt.subplots()
        ax.bar(feats, fi)
        ax.set_title('Важность признаков XGBoost')
        ax.set_xlabel('Признаки'); ax.set_ylabel('Важность'); plt.xticks(rotation=45, ha='right')
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            fig.savefig(tmp.name, bbox_inches='tight'); feat_imp = tmp.name
        plt.close(fig)
        doc.add_picture(feat_imp, width=Inches(6))
        doc.add_heading('Диаграмма важности признаков XGBoost', level=3)
        doc.add_paragraph('Показывает вклад каждого признака в итоговый прогноз модели XGBoost.', style='Intense Quote')

    doc.add_heading('Рекомендации', level=2)
    doc.add_paragraph('Для быстрого получения точечных оценок используйте XGBoost.', style='List Bullet')
    doc.add_paragraph('Для анализа трендов и сезонности обращайтесь к LSTM.', style='List Bullet')
    doc.add_paragraph('Комбинируйте оба прогноза и учитывайте их метрики при принятии решений.', style='List Bullet')

    doc.save(report_path)
    logger.info('Report generation completed', extra={'path': str(report_path)})
